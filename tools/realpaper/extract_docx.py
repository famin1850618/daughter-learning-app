#!/usr/bin/env python3
"""extract_docx.py — V3.12.21 起 docx 入库主工具（深圳真题源专用）

输入: <docx_path>
输出 cache/docx/<sha1>/:
  - raw.txt        libreoffice 转纯文本（fallback python-docx）
  - media/         word/media/ 提取图（WMF/EMF 转 PNG，大图压缩）
  - media_meta.json  每张图的:
      original_name / format / size_bytes / size_kb / image_data_b64 / hash
  - paragraph_image_map.json  段落索引 → 图片名 list
  - structure.json  (题号 + 段落 ranges + 图片归属推断)

不入库 batch JSON——让 4a annotate 阶段做（agent 用 raw.txt + media + paragraph_image_map）

Why docx vs PDF (V3.12.20.1 痛点对比):
  - docx 文字层完美（OOXML 标准，无私有字体 cid 编码）
  - 拼音保留声调（§1.5 原题保真天然兼容）
  - 内嵌图原图直接拿（不需 pdftoppm 截图，质量翻倍）
  - 图位置 XML 上下文可推断归属哪道题（比 PDF 截图更精准）

V3.12.21 已知边界:
  - 数学 docx 59% 图是 WMF 矢量（Word 几何图默认 WMF）→ libreoffice 转 PNG（1200x900）
  - 图 99% < 200KB 不需压缩；> 500KB 才走 PIL 压缩
  - 老 doc（非 docx）需先 libreoffice --convert-to docx 转换
"""
import argparse
import base64
import hashlib
import io
import json
import os
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Optional

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import docx as docx_lib  # python-docx
except ImportError:
    docx_lib = None


CACHE_ROOT = Path.home() / 'daughter_learning_app' / '.cache' / 'docx'
CACHE_ROOT.mkdir(parents=True, exist_ok=True)

MAX_IMAGE_KB = 500  # § 9.2 V3.12.21 修订：放宽到 500KB（深圳源 max 670KB 也覆盖）


# ----------------------------------------------------------------------
# Step 1: extract text
# ----------------------------------------------------------------------

def extract_text_libreoffice(docx_path: Path, out_dir: Path) -> Optional[str]:
    """用 libreoffice headless 转 txt（首选）"""
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'txt:Text (encoded):UTF8',
             '--outdir', str(out_dir), str(docx_path)],
            capture_output=True, text=True, timeout=120
        )
        txt_path = out_dir / (docx_path.stem + '.txt')
        if txt_path.exists():
            return txt_path.read_text(encoding='utf-8')
    except Exception as e:
        print(f'  ⚠ libreoffice extract failed: {e}', file=sys.stderr)
    return None


def extract_text_pythondocx(docx_path: Path) -> Optional[str]:
    """python-docx 路径（fallback）"""
    if docx_lib is None:
        return None
    try:
        d = docx_lib.Document(str(docx_path))
        lines = []
        for p in d.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        lines.append(t)
        return '\n'.join(lines)
    except Exception as e:
        print(f'  ⚠ python-docx extract failed: {e}', file=sys.stderr)
        return None


def extract_text_paragraph_aligned(docx_path: Path) -> Optional[str]:
    """V3.12.22 A1: 用 python-docx 按 word/document.xml 段顺序输出 raw.txt
    每行 = 一个 docx 段落（含空段保留）→ raw.txt 行号 == paragraph_image_map 段索引

    解决 V3.12.21 痛点：libreoffice 转 txt 拼接段落后，raw.txt 行号与 pmap 段索引不对应。
    """
    if docx_lib is None:
        return None
    try:
        d = docx_lib.Document(str(docx_path))
        # 按 document.xml 顺序遍历，包括空段（与 paragraph_image_map 索引同步）
        lines = []
        # python-docx 的 d.paragraphs 是 body level 段落（不含 table 内部）
        # paragraph_image_map 解析的也是 body 段落（同源），所以索引天然对齐
        for p in d.paragraphs:
            t = p.text  # 保留空段（不 strip 整行）
            lines.append(t)
        return '\n'.join(lines)
    except Exception as e:
        print(f'  ⚠ paragraph-aligned extract failed: {e}', file=sys.stderr)
        return None


# ----------------------------------------------------------------------
# Step 2: extract media
# ----------------------------------------------------------------------

def is_raster_format(name: str) -> bool:
    return name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))


def is_vector_format(name: str) -> bool:
    return name.lower().endswith(('.wmf', '.emf', '.svg'))


def convert_wmf_to_png(wmf_bytes: bytes, target_w: int = 1200) -> Optional[bytes]:
    """用 libreoffice 把 WMF/EMF 转 PNG（1200 像素宽）"""
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        wmf_path = td / 'tmp.wmf'
        wmf_path.write_bytes(wmf_bytes)
        try:
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'png',
                 '--outdir', str(td), str(wmf_path)],
                capture_output=True, timeout=60
            )
            png_path = td / 'tmp.png'
            if png_path.exists():
                return png_path.read_bytes()
        except Exception:
            pass
    return None


def compress_if_needed(img_bytes: bytes, max_kb: int = MAX_IMAGE_KB) -> bytes:
    """如果图 > max_kb，PIL resize + 重压缩到 max_kb 内"""
    if len(img_bytes) <= max_kb * 1024:
        return img_bytes
    if Image is None:
        return img_bytes  # 没 PIL 就不压（带原大小通过）
    try:
        img = Image.open(io.BytesIO(img_bytes))
        # 如果 PNG 直接转 JPEG 大幅压缩（除非透明）
        has_alpha = img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info)
        # 先按比例缩放最长边到 1600
        if max(img.size) > 1600:
            img.thumbnail((1600, 1600), Image.LANCZOS)
        # 反复降质压缩到 max_kb 内
        buf = io.BytesIO()
        if has_alpha:
            img.save(buf, 'PNG', optimize=True)
        else:
            img.convert('RGB').save(buf, 'JPEG', quality=85, optimize=True)
        out = buf.getvalue()
        if len(out) <= max_kb * 1024:
            return out
        # 还是大，进一步降质
        for q in [75, 65, 55, 45]:
            buf = io.BytesIO()
            img.convert('RGB').save(buf, 'JPEG', quality=q, optimize=True)
            out = buf.getvalue()
            if len(out) <= max_kb * 1024:
                return out
        return out
    except Exception as e:
        print(f'  ⚠ compress failed: {e}', file=sys.stderr)
        return img_bytes


def extract_media(docx_path: Path, out_dir: Path) -> list:
    """提取 word/media/* 所有图。WMF/EMF 转 PNG，大图压缩。返回 meta list."""
    media_dir = out_dir / 'media'
    media_dir.mkdir(exist_ok=True)
    metas = []

    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if not name.startswith('word/media/'):
                continue
            base_name = os.path.basename(name)
            if not base_name:
                continue
            try:
                raw = z.read(name)
            except Exception:
                continue

            original_format = base_name.rsplit('.', 1)[-1].lower() if '.' in base_name else 'bin'
            output_name = base_name
            output_bytes = raw

            # WMF/EMF 转 PNG
            if is_vector_format(base_name) and not base_name.endswith('.svg'):
                converted = convert_wmf_to_png(raw)
                if converted:
                    output_bytes = converted
                    output_name = base_name.rsplit('.', 1)[0] + '.png'
                    print(f'  · {base_name} ({original_format} {len(raw)//1024}KB) → PNG {len(output_bytes)//1024}KB')
                else:
                    print(f'  ⚠ {base_name} WMF→PNG 转换失败，保留原文件')

            # 大图压缩
            if is_raster_format(output_name) and len(output_bytes) > MAX_IMAGE_KB * 1024:
                before_kb = len(output_bytes) // 1024
                output_bytes = compress_if_needed(output_bytes)
                after_kb = len(output_bytes) // 1024
                if before_kb != after_kb:
                    print(f'  · {output_name} 压缩 {before_kb}KB → {after_kb}KB')

            # 写文件 + 计算 base64
            (media_dir / output_name).write_bytes(output_bytes)
            b64 = base64.b64encode(output_bytes).decode()
            mime = 'image/png' if output_name.endswith('.png') else \
                   'image/jpeg' if output_name.endswith(('.jpg', '.jpeg')) else \
                   'image/svg+xml' if output_name.endswith('.svg') else 'image/png'

            metas.append({
                'original_name': base_name,
                'original_format': original_format,
                'output_name': output_name,
                'output_format': output_name.rsplit('.', 1)[-1].lower(),
                'size_bytes': len(output_bytes),
                'size_kb': len(output_bytes) // 1024,
                'data_uri': f'data:{mime};base64,{b64}',
                'hash': hashlib.sha1(output_bytes).hexdigest()[:12],
            })

    return metas


# ----------------------------------------------------------------------
# Step 3: paragraph → image owner mapping
# ----------------------------------------------------------------------

NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
}

M_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'


# ----------------------------------------------------------------------
# OMath → LaTeX (V3.12.22 Issue 1: Famin 决策升级工具救公式段)
# ----------------------------------------------------------------------

# 常见运算符映射（OMath 文本 → LaTeX）
OMATH_OPERATOR_MAP = {
    '×': r'\times ',
    '÷': r'\div ',
    '·': r'\cdot ',
    '−': '-', '–': '-', '—': '-',  # 各种 dash 统一 -
    '≤': r'\le ', '≥': r'\ge ',
    '≠': r'\ne ', '≈': r'\approx ',
    '∞': r'\infty ',
    'π': r'\pi ',
    '°': r'^{\circ}',
}


def _omath_text_escape(s: str) -> str:
    """OMath 文本节点转 LaTeX 安全字符（运算符替换）"""
    if not s:
        return ''
    out = []
    for ch in s:
        if ch in OMATH_OPERATOR_MAP:
            out.append(OMATH_OPERATOR_MAP[ch])
        elif ch in '%&_#$^{}~\\':
            out.append('\\' + ch)
        else:
            out.append(ch)
    return ''.join(out)


def _omath_child_text(elem, tag_local: str) -> str:
    """递归获取子节点 m:tag 的 LaTeX 内容（无该子节点返回空串）"""
    child = elem.find(M_NS + tag_local)
    if child is None:
        return ''
    return _omath_to_latex_recursive(child)


def _omath_to_latex_recursive(elem) -> str:
    """递归转 OMath XML 元素 → LaTeX 字符串"""
    if elem is None:
        return ''
    # 提取 namespace 后的 local tag
    tag = elem.tag
    if tag.startswith(M_NS):
        local = tag[len(M_NS):]
    else:
        local = tag.rsplit('}', 1)[-1] if '}' in tag else tag

    # 1. 文本节点 m:t → 直接输出 escape 后文字
    if local == 't':
        return _omath_text_escape(elem.text or '')

    # 2. m:r run → 递归子节点 m:t
    if local == 'r':
        parts = []
        for child in elem:
            if child.tag == M_NS + 't':
                parts.append(_omath_text_escape(child.text or ''))
        return ''.join(parts)

    # 3. m:f 分数 → \frac{num}{den}
    if local == 'f':
        num = _omath_child_text(elem, 'num')
        den = _omath_child_text(elem, 'den')
        return f'\\frac{{{num}}}{{{den}}}'

    # 4. m:sSup 上标 → {base}^{sup}
    if local == 'sSup':
        base = _omath_child_text(elem, 'e')
        sup = _omath_child_text(elem, 'sup')
        return f'{{{base}}}^{{{sup}}}'

    # 5. m:sSub 下标 → {base}_{sub}
    if local == 'sSub':
        base = _omath_child_text(elem, 'e')
        sub = _omath_child_text(elem, 'sub')
        return f'{{{base}}}_{{{sub}}}'

    # 6. m:sSubSup 上下标 → {base}_{sub}^{sup}
    if local == 'sSubSup':
        base = _omath_child_text(elem, 'e')
        sub = _omath_child_text(elem, 'sub')
        sup = _omath_child_text(elem, 'sup')
        return f'{{{base}}}_{{{sub}}}^{{{sup}}}'

    # 7. m:rad 根号
    if local == 'rad':
        deg = _omath_child_text(elem, 'deg')
        e = _omath_child_text(elem, 'e')
        if deg:
            return f'\\sqrt[{deg}]{{{e}}}'
        return f'\\sqrt{{{e}}}'

    # 8. m:nary n 元运算（求和/积分）
    if local == 'nary':
        # 找 m:naryPr/m:chr 取运算符（默认 \int）
        naryPr = elem.find(M_NS + 'naryPr')
        op = '\\int '
        if naryPr is not None:
            chr_elem = naryPr.find(M_NS + 'chr')
            if chr_elem is not None:
                v = chr_elem.get(M_NS.replace('{','').replace('}','') and '{http://schemas.openxmlformats.org/officeDocument/2006/math}val', '')
                # 上面 v 取法不对，简化：直接 attrib 取
                v = chr_elem.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', '')
                if v == '∑':
                    op = '\\sum '
                elif v == '∏':
                    op = '\\prod '
                elif v == '∫':
                    op = '\\int '
                elif v:
                    op = _omath_text_escape(v)
        sub = _omath_child_text(elem, 'sub')
        sup = _omath_child_text(elem, 'sup')
        e = _omath_child_text(elem, 'e')
        out = op
        if sub:
            out += f'_{{{sub}}}'
        if sup:
            out += f'^{{{sup}}}'
        out += f'{{{e}}}'
        return out

    # 9. m:d 括号
    if local == 'd':
        # 默认圆括号；m:dPr/m:begChr / m:endChr 改括号
        dPr = elem.find(M_NS + 'dPr')
        beg, end = '(', ')'
        if dPr is not None:
            begChr = dPr.find(M_NS + 'begChr')
            endChr = dPr.find(M_NS + 'endChr')
            if begChr is not None:
                v = begChr.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', '(')
                beg = v
            if endChr is not None:
                v = endChr.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', ')')
                end = v
        # m:e 可能多个（逗号分隔）
        es = [_omath_to_latex_recursive(e) for e in elem.findall(M_NS + 'e')]
        body = ','.join(es)
        # LaTeX 中 { } 需 \{ \}
        if beg == '{':
            beg = r'\{'
        if end == '}':
            end = r'\}'
        return f'{beg}{body}{end}'

    # 10. m:func 函数
    if local == 'func':
        fname = _omath_child_text(elem, 'fName')
        e = _omath_child_text(elem, 'e')
        return f'{fname}({e})'

    # 11. m:bar 上下划线
    if local == 'bar':
        e = _omath_child_text(elem, 'e')
        return f'\\overline{{{e}}}'

    # 12. m:m 矩阵 → 用 pmatrix
    if local == 'm':
        rows = []
        for mr in elem.findall(M_NS + 'mr'):
            cells = [_omath_to_latex_recursive(e) for e in mr.findall(M_NS + 'e')]
            rows.append(' & '.join(cells))
        return '\\begin{pmatrix}' + ' \\\\ '.join(rows) + '\\end{pmatrix}'

    # 默认：递归子节点拼接（覆盖 m:e, m:num, m:den, m:sub, m:sup 等容器）
    parts = []
    for child in elem:
        parts.append(_omath_to_latex_recursive(child))
    return ''.join(parts)


def omath_to_latex(omath_elem) -> str:
    """OMath XML 节点 → LaTeX 字符串（顶层包裹 $...$）"""
    body = _omath_to_latex_recursive(omath_elem)
    return f'${body}$' if body else ''


def extract_omath_per_paragraph(docx_path: Path) -> dict:
    """解析 word/document.xml 找每个段落里的 OMath 公式 → LaTeX。
    返回 {paragraph_index: [latex_str, ...]}（每段可能多个公式）

    V3.12.22 Issue 1 (Famin 决策 D)：升级工具救 OMath 段，避免整题跳。
    """
    import xml.etree.ElementTree as ET

    para_to_latex = {}
    try:
        with zipfile.ZipFile(docx_path) as z:
            doc_xml = z.read('word/document.xml')
        root = ET.fromstring(doc_xml)
        body = root.find('w:body', NS)
        if body is None:
            return para_to_latex

        for idx, p in enumerate(body.findall('w:p', NS)):
            latex_list = []
            # 找该段所有 oMath 节点（含 oMathPara 包裹的）
            for omath in p.iter(M_NS + 'oMath'):
                tex = omath_to_latex(omath)
                if tex:
                    latex_list.append(tex)
            if latex_list:
                # 统一用 str key（cached JSON load 时 key 自然是 str，新生成时也用 str）
                para_to_latex[str(idx)] = latex_list
    except Exception as e:
        print(f'  ⚠ parse omath failed: {e}', file=sys.stderr)

    return para_to_latex


def parse_paragraph_image_map(docx_path: Path) -> dict:
    """从 word/document.xml 解析每个 paragraph 包含哪些 image rId。
    返回 {paragraph_index: [image_basename, ...]}
    需要 word/_rels/document.xml.rels 把 rId 映射到 word/media/* 文件名"""
    import xml.etree.ElementTree as ET

    para_to_images = {}
    rid_to_target = {}

    with zipfile.ZipFile(docx_path) as z:
        # 1. 解 _rels 拿 rId → 文件名
        try:
            rels_xml = z.read('word/_rels/document.xml.rels')
            rels_root = ET.fromstring(rels_xml)
            for rel in rels_root.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                rid = rel.get('Id')
                target = rel.get('Target')
                if rid and target and 'media/' in target:
                    # target 形如 'media/image1.png'
                    rid_to_target[rid] = os.path.basename(target)
        except Exception as e:
            print(f'  ⚠ parse rels failed: {e}', file=sys.stderr)

        # 2. 解 document.xml 找每个 <w:p> 里的 <pic:pic> 或 <a:blip>
        try:
            doc_xml = z.read('word/document.xml')
            root = ET.fromstring(doc_xml)
            body = root.find('w:body', NS)
            if body is None:
                return para_to_images
            paragraphs = body.findall('w:p', NS)
            for idx, p in enumerate(paragraphs):
                images = []
                # 找 <a:blip r:embed="rIdN"/> 元素
                for blip in p.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                    rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rid and rid in rid_to_target:
                        images.append(rid_to_target[rid])
                # 也找 <v:imagedata r:id="rIdN"/>（VML 老格式）
                for imgdata in p.iter('{urn:schemas-microsoft-com:vml}imagedata'):
                    rid = imgdata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if rid and rid in rid_to_target:
                        images.append(rid_to_target[rid])
                if images:
                    para_to_images[idx] = images
        except Exception as e:
            print(f'  ⚠ parse document.xml failed: {e}', file=sys.stderr)

    return para_to_images


# ----------------------------------------------------------------------
# Step 4: question number recognition (heuristic)
# ----------------------------------------------------------------------

QNUM_PATTERNS = [
    re.compile(r'^([一二三四五六七八九十]+)[、．.]\s*'),    # 一、 一． 一.
    re.compile(r'^(\d+)\s*[．.]\s*'),                       # 1． 1.
    re.compile(r'^（(\d+)）'),                              # （1）
    re.compile(r'^\((\d+)\)'),                             # (1)
    re.compile(r'^[(（]([一二三四五六七八九十]+)[)）]'),    # （一）
]


def detect_question_anchors(text: str) -> list:
    """简单题号识别：返回每行的"是否题号""题号文本"映射"""
    lines = text.split('\n')
    anchors = []
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        for pat in QNUM_PATTERNS:
            m = pat.match(line_stripped)
            if m:
                anchors.append({
                    'line_index': i,
                    'qnum_text': m.group(1),
                    'pattern': pat.pattern[:30],
                    'line_preview': line_stripped[:80],
                })
                break
    return anchors


# ----------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------

def process_docx(docx_path: Path, force: bool = False) -> dict:
    """主流程"""
    if not docx_path.exists():
        return {'error': f'file not found: {docx_path}'}
    sha1 = hashlib.sha1(docx_path.read_bytes()[:65536] + str(docx_path.stat().st_size).encode()).hexdigest()[:16]
    out_dir = CACHE_ROOT / sha1
    out_dir.mkdir(exist_ok=True, parents=True)

    print(f'[extract_docx] {docx_path.name}')
    print(f'  → cache: {out_dir}')

    # 1. extract text
    raw_path = out_dir / 'raw.txt'
    if raw_path.exists() and not force:
        text = raw_path.read_text(encoding='utf-8')
        print(f'  · text 已缓存 ({len(text)} chars)')
    else:
        text = extract_text_libreoffice(docx_path, out_dir)
        if text is None or len(text) < 100:
            print('  · libreoffice 失败/输出过短，fallback python-docx')
            text = extract_text_pythondocx(docx_path) or ''
        raw_path.write_text(text, encoding='utf-8')
        print(f'  · text 提取 {len(text)} chars')

    # 1.5. V3.12.22 A1: 段对齐 raw_aligned.txt（行号 == paragraph_image_map 段索引）
    # 修 V3.12.21 paragraph_image_map 索引与 raw.txt 行号不对应的痛点
    aligned_path = out_dir / 'raw_aligned.txt'
    if aligned_path.exists() and not force:
        print(f'  · raw_aligned 已缓存')
    else:
        aligned = extract_text_paragraph_aligned(docx_path)
        if aligned is not None:
            aligned_path.write_text(aligned, encoding='utf-8')
            n_para = len(aligned.split('\n'))
            print(f'  · raw_aligned 输出 {n_para} 段（与 paragraph_image_map 索引同步）')
        else:
            print(f'  ⚠ raw_aligned 提取失败（pmap 推断图归属时 fallback 用 raw.txt）')

    # 2. extract media
    media_meta_path = out_dir / 'media_meta.json'
    if media_meta_path.exists() and not force:
        media_metas = json.loads(media_meta_path.read_text(encoding='utf-8'))
        print(f'  · media 已缓存 ({len(media_metas)} 张)')
    else:
        print(f'  · 提取 media...')
        media_metas = extract_media(docx_path, out_dir)
        # 不存 base64 到磁盘 cache（base64 太大），只存路径 + meta
        meta_lite = []
        for m in media_metas:
            mc = {k: v for k, v in m.items() if k != 'data_uri'}
            mc['data_uri_size'] = len(m['data_uri'])
            meta_lite.append(mc)
        media_meta_path.write_text(json.dumps(meta_lite, ensure_ascii=False, indent=2), encoding='utf-8')
        print(f'  · media 提取 {len(media_metas)} 张')

    # 3. paragraph → image map
    pmap_path = out_dir / 'paragraph_image_map.json'
    if pmap_path.exists() and not force:
        para_image_map = json.loads(pmap_path.read_text(encoding='utf-8'))
        print(f'  · pmap 已缓存 ({len(para_image_map)} 段含图)')
    else:
        para_image_map = parse_paragraph_image_map(docx_path)
        pmap_path.write_text(json.dumps(para_image_map, ensure_ascii=False, indent=2), encoding='utf-8')
        print(f'  · pmap 解析 {len(para_image_map)} 段含图')

    # 3.5. V3.12.22 Issue 1: OMath 公式段 → LaTeX
    # docx 用 Word OMath 嵌入数学公式（XML 结构化，非位图），libreoffice 转 txt 时丢失。
    # 本步骤直接解析 word/document.xml 的 m:oMath 节点 → LaTeX，输出 omath_map.json。
    # annotate worker 用 omath_map[段索引] 替换 raw_aligned 对应段的空白为 LaTeX 公式。
    omath_path = out_dir / 'omath_map.json'
    if omath_path.exists() and not force:
        omath_map = json.loads(omath_path.read_text(encoding='utf-8'))
        print(f'  · omath 已缓存 ({len(omath_map)} 段含公式)')
    else:
        omath_map = extract_omath_per_paragraph(docx_path)
        omath_path.write_text(json.dumps(omath_map, ensure_ascii=False, indent=2), encoding='utf-8')
        total_formulas = sum(len(v) for v in omath_map.values())
        print(f'  · omath 解析 {len(omath_map)} 段含公式（共 {total_formulas} 个 OMath → LaTeX）')

    # 3.6. V3.12.22 Issue 1: 合成 raw_with_omath.txt
    # 在 raw_aligned 基础上，每段后面（如有公式）追加 [OMATH: $latex$] 标记。
    # annotate worker 读这个文件可直接复制 LaTeX 入 content，无需另查 omath_map。
    aligned_omath_path = out_dir / 'raw_with_omath.txt'
    if aligned_omath_path.exists() and not force:
        print(f'  · raw_with_omath 已缓存')
    elif aligned_path.exists():
        aligned_lines = aligned_path.read_text(encoding='utf-8').split('\n')
        merged = []
        for idx, line in enumerate(aligned_lines):
            merged.append(line)
            formulas = omath_map.get(str(idx), [])
            if formulas:
                merged.append(f'  [OMATH: {" | ".join(formulas)}]')
        aligned_omath_path.write_text('\n'.join(merged), encoding='utf-8')
        print(f'  · raw_with_omath 输出（{len(omath_map)} 段注入 OMath LaTeX 标记）')

    # 4. question anchors
    anchors = detect_question_anchors(text)
    print(f'  · 识别 {len(anchors)} 个题号锚点')

    # 5. structure summary
    structure = {
        'sha1': sha1,
        'docx_path': str(docx_path),
        'text_length': len(text),
        'media_count': len(media_metas) if isinstance(media_metas, list) else 0,
        'para_with_images': len(para_image_map),
        'para_with_omath': len(omath_map),
        'omath_total_formulas': sum(len(v) for v in omath_map.values()),
        'qnum_anchors': anchors,
        'media_summary': [
            {
                'output_name': m.get('output_name'),
                'output_format': m.get('output_format'),
                'size_kb': m.get('size_kb'),
                'original_format': m.get('original_format'),
            } for m in (media_metas if isinstance(media_metas, list) else [])
        ],
    }
    (out_dir / 'structure.json').write_text(
        json.dumps(structure, ensure_ascii=False, indent=2), encoding='utf-8'
    )
    return structure


def main():
    p = argparse.ArgumentParser(description='V3.12.21 docx 入库主工具')
    p.add_argument('docx_path', help='Path to .docx file')
    p.add_argument('--force', action='store_true', help='强制重新提取（忽略 cache）')
    p.add_argument('--summary', action='store_true', help='只打印 summary 不执行')
    args = p.parse_args()

    docx_path = Path(args.docx_path).expanduser()
    if args.summary:
        sha1 = hashlib.sha1(docx_path.read_bytes()[:65536] + str(docx_path.stat().st_size).encode()).hexdigest()[:16]
        out_dir = CACHE_ROOT / sha1
        sp = out_dir / 'structure.json'
        if sp.exists():
            print(sp.read_text(encoding='utf-8'))
        else:
            print(f'no cache for {docx_path}', file=sys.stderr)
        return

    structure = process_docx(docx_path, force=args.force)
    print()
    print(json.dumps({
        'sha1': structure.get('sha1'),
        'text_length': structure.get('text_length'),
        'media_count': structure.get('media_count'),
        'para_with_images': structure.get('para_with_images'),
        'qnum_anchors': len(structure.get('qnum_anchors', [])),
    }, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
